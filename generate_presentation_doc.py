"""
Génère le document Word de présentation du projet Urban Data Explorer (EFREI).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles de base ─────────────────────────────────────────────────────────────

def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x10, 0x3F, 0x6E)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x10, 0x6E, 0x45)
    return p

def heading3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # En-tête
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '103F6E')
        tcPr.append(shd)
    # Lignes
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    return table

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Urban Data Explorer")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x10, 0x3F, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Plateforme d'analyse urbaine et immobilière — Paris")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x10, 0x6E, 0x45)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info.add_run("Projet EFREI  •  Présentation jury  •  2024–2025")
run3.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = tagline.add_run(
    "Révéler les inégalités cachées du territoire parisien\n"
    "à travers des données ouvertes, des algorithmes géospatiaux\n"
    "et une carte interactive en temps réel."
)
run4.font.size = Pt(13)
run4.italic = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════════════════════

heading1("Table des matières")
toc_items = [
    ("1", "Vision du projet et objectifs"),
    ("2", "Architecture globale — vue d'ensemble"),
    ("3", "Les données sources (couche Bronze)"),
    ("4", "Pipeline de transformation (Bronze → Silver → Gold)"),
    ("5", "Les 4 indicateurs personnalisés — algorithmes détaillés"),
    ("   5.1", "Indice de Vivabilité Familiale (composite)"),
    ("   5.2", "Accessibilité aux Transports"),
    ("   5.3", "Confort Thermique Urbain"),
    ("   5.4", "Logement — Prix de vente et loyers"),
    ("6", "Base de données — MySQL (Gold) et MongoDB (Analytics)"),
    ("7", "API REST — FastAPI"),
    ("8", "Frontend — Carte interactive (Next.js + Mapbox GL)"),
    ("9", "Système d'authentification JWT"),
    ("10", "Analytics utilisateurs — MongoDB"),
    ("11", "Stack technique complète"),
    ("12", "Ce qui fonctionne bien"),
    ("13", "Limites actuelles et axes d'amélioration"),
    ("14", "Conclusion"),
]
for num, title_text in toc_items:
    p = doc.add_paragraph()
    run_n = p.add_run(f"{num}  ")
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_t = p.add_run(title_text)
    run_t.font.size = Pt(11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. VISION DU PROJET
# ══════════════════════════════════════════════════════════════════════════════

heading1("1. Vision du projet et objectifs")

body(
    "Urban Data Explorer est une plateforme de données géospatiales dédiée à Paris. "
    "L'objectif est de transformer des données publiques brutes (recensements INSEE, "
    "transactions immobilières DVF, établissements BDCOM, hôpitaux, transports, espaces "
    "verts, écoles…) en indicateurs synthétiques visualisables sur une carte interactive, "
    "accessible après authentification."
)

heading2("Problématique centrale")
body(
    "Comment identifier objectivement les zones parisiennes les plus favorables à la vie "
    "familiale, en combinant accessibilité aux services, qualité de l'environnement et "
    "coût du logement, à partir de données 100 % ouvertes ?"
)

heading2("Ce que le projet produit")
bullet("Une carte choroplèthe interactive à l'échelle IRIS (800+ zones) et arrondissement (20 zones)")
bullet("4 indicateurs composites sur 0–10, recalculables en temps réel par l'utilisateur")
bullet("Un moteur de recherche et de classement des zones selon les critères choisis")
bullet("Un système d'authentification sécurisé (JWT + bcrypt)")
bullet("Un pipeline de données reproductible Bronze → Silver → Gold")

heading2("Données clés")
add_table(
    ["Dimension", "Valeur"],
    [
        ["Zones IRIS analysées", "~992 (Paris)"],
        ["Sources de données distinctes", "12 jeux de données publics"],
        ["Indicateurs composites", "4 (vivabilité, transport, thermique, logement)"],
        ["Piliers dans la vivabilité", "5 (écoles, santé, transport, services, espaces verts)"],
        ["Arrondissements couverts", "20"],
        ["Technologies back-end", "Python / FastAPI / MySQL / MongoDB"],
        ["Technologies front-end", "Next.js / React 19 / Mapbox GL / Tailwind CSS"],
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE GLOBALE
# ══════════════════════════════════════════════════════════════════════════════

heading1("2. Architecture globale — vue d'ensemble")

body(
    "Le projet suit le patron d'architecture « Medallion » (Médaillon), standard dans "
    "l'ingénierie des données modernes. Chaque couche a une responsabilité unique et "
    "les données ne remontent jamais en arrière."
)

heading2("Schéma d'architecture")
code_block(
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  SOURCES EXTERNES (Internet / Google Drive)                    │\n"
    "│  INSEE · DVF · BDCOM · OpenData Paris · RATP · IGN            │\n"
    "└────────────────────────────┬────────────────────────────────────┘\n"
    "                             │ téléchargement (gdown)\n"
    "                             ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  COUCHE BRONZE  (data/bronze/)                                  │\n"
    "│  Données brutes, non modifiées, formats originaux              │\n"
    "│  .xlsx · .csv · .geojson · .pdf · .kml · .txt                 │\n"
    "└────────────────────────────┬────────────────────────────────────┘\n"
    "                             │ nettoyage / filtre / normalisation\n"
    "                             ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  COUCHE SILVER  (data/silver/)                                  │\n"
    "│  CSV propres, coordonnées validées, colonnes harmonisées       │\n"
    "│  Pas de calcul métier — seulement de la qualité de données     │\n"
    "└────────────────────────────┬────────────────────────────────────┘\n"
    "                             │ calculs géospatiaux + scoring\n"
    "                             ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  COUCHE GOLD  (data/gold/ + MySQL)                              │\n"
    "│  Indicateurs 0–10 par zone IRIS, agrégats arrondissement       │\n"
    "│  CSV + Parquet + tables MySQL (requêtes API en O(1))           │\n"
    "└────────────────────────────┬────────────────────────────────────┘\n"
    "                             │ REST API\n"
    "                             ▼\n"
    "┌───────────────────────┐   ┌─────────────────────────────────────┐\n"
    "│  FASTAPI  (:8000)     │   │  MONGODB  (Atlas cloud)              │\n"
    "│  /map  /indicators    │   │  Authentification · Analytics       │\n"
    "│  /auth  /analytics    │   │  Collections: users, zone_clicks    │\n"
    "└───────────────────────┘   └─────────────────────────────────────┘\n"
    "                             │\n"
    "                             ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  FRONTEND  Next.js  (:3000)                                     │\n"
    "│  Landing page → /login → /dashboard (carte Mapbox)             │\n"
    "│  Carte choroplèthe IRIS/arrondissement · Filtres · Classement  │\n"
    "└─────────────────────────────────────────────────────────────────┘"
)

heading2("Flux de données en pratique")
body(
    "Lors du démarrage de l'API (uvicorn), FastAPI charge en mémoire tous les fichiers "
    "Gold (CSV, Parquet, GeoJSON) via le DataStore. Cette stratégie évite toute requête "
    "SQL pendant la durée de vie des requêtes HTTP — les réponses sont quasi-instantanées. "
    "La base MySQL est donc uniquement utilisée comme stockage persistant, pas comme "
    "moteur de requêtes temps réel."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. DONNÉES SOURCES (BRONZE)
# ══════════════════════════════════════════════════════════════════════════════

heading1("3. Les données sources (couche Bronze)")

body(
    "Toutes les données utilisées sont publiques et librement accessibles. "
    "Voici les 12 sources principales avec leur rôle dans le projet :"
)

add_table(
    ["Source", "Format", "Contenu", "Utilisation dans le projet"],
    [
        ["INSEE — Recensement 2022", "CSV", "Population par zone IRIS, catégories socio-professionnelles", "Base de pondération pour tous les scores par habitant"],
        ["IGN — Contours IRIS", "GeoJSON + XLSX", "Polygones géographiques des 992 zones IRIS parisiennes", "Fond de carte, jointures spatiales"],
        ["Mairie de Paris — Arrondissements", "GeoJSON", "Polygones des 20 arrondissements", "Agrégation à l'échelle supérieure"],
        ["DVF — Demandes de Valeurs Foncières", "TXT (séparateur |)", "Transactions immobilières 2025 : type, surface, prix", "Indicateur prix de vente au m²"],
        ["BDCOM 2023 — Base commerciale", "CSV + XLSX OD", "5 000+ établissements classifiés (niv8 : alimentaire, santé, culturel…)", "Services quotidiens, santé, commerces"],
        ["Hôpitaux franciliens", "CSV", "Coordonnées GPS des hôpitaux publics", "Pilier santé de la vivabilité"],
        ["Écoles publiques parisiennes", "3 × XLSX", "Collèges, maternelles, élémentaires avec coordonnées", "Pilier écoles de la vivabilité"],
        ["Espaces verts Paris", "GeoJSON", "Parcs, jardins, squares avec surface réelle (m²)", "Pilier espaces verts de la vivabilité"],
        ["IDFM — Arrêts de transport", "CSV", "Arrêts métro, RER, bus, tram, câbleway", "Pilier transport de la vivabilité"],
        ["Vélib' — Stations", "CSV", "Stations Vélib' avec coordonnées GPS", "Composante vélo du score transport"],
        ["Observatoire des Loyers — Paris", "CSV + KML + XLSX", "Loyers médians 2024 par zone élémentaire", "Indicateur loyer par arrondissement"],
        ["Historique prix immobiliers", "PDF", "Prix médians au m² par arrondissement (série historique)", "Indicateur prix de vente par arrondissement"],
        ["Îlots de fraîcheur Paris", "GeoJSON", "Zones de fraîcheur urbaine (parcs frais)", "Composante fraîcheur du confort thermique"],
        ["Arbres de Paris", "GeoJSON", "Inventaire des arbres avec coordonnées", "Composante densité arborée du confort thermique"],
    ]
)

heading2("Filtres appliqués dès la couche Silver")
body("Avant tout calcul, les données brutes sont nettoyées selon des règles métier strictes :")
bullet("DVF : conservation uniquement des Appartements et Maisons, surface entre 5 et 1 000 m², prix entre 1 000 et 50 000 €/m²")
bullet("IRIS : uniquement les zones résidentielles (TYP_IRIS = 'H') de Paris (code commune commençant par '75')")
bullet("Écoles : 4 types retenus — Collège, Maternelle, Élémentaire, Polyvalent")
bullet("Transport : harmonisation des libellés de type (metro, rail, tram, bus, cableway, velib)")
bullet("Espaces verts : validation des géométries, calcul de la surface_totale_reelle")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PIPELINE DE TRANSFORMATION
# ══════════════════════════════════════════════════════════════════════════════

heading1("4. Pipeline de transformation (Bronze → Silver → Gold)")

heading2("4.1 Orchestration — run_pipeline.py")
body(
    "Le fichier run_pipeline.py est le point d'entrée unique. Il accepte des arguments "
    "en ligne de commande pour exécuter tout ou partie du pipeline :"
)
code_block(
    "python run_pipeline.py           # Pipeline complet\n"
    "python run_pipeline.py --silver  # Nettoyage uniquement\n"
    "python run_pipeline.py --gold    # Calcul des indicateurs uniquement\n"
    "python run_pipeline.py --bronze  # Téléchargement uniquement"
)

heading2("4.2 Couche Silver — scripts de nettoyage")
add_table(
    ["Script", "Entrée (Bronze)", "Sortie (Silver)", "Opérations clés"],
    [
        ["population.py", "base-ic-evol-struct-pop-2022.CSV", "population_paris.csv", "Filtrage Paris + zones résidentielles, renommage colonnes"],
        ["iris.py", "iris.xlsx + iris.geojson", "iris_paris.csv", "Extraction métadonnées IRIS, jointure noms zones"],
        ["schools.py", "3 fichiers XLSX", "schools_merged.csv", "Fusion des 3 jeux, déduplication, validation lat/lng"],
        ["dvf.py", "ValeursFoncieres-2025.txt", "dvf_paris_clean.csv", "Filtre type/surface/prix, calcul prix_m2"],
        ["bdcom.py", "BDCOM_2023.csv + OD XLSX", "bdcom_paris_clean.csv", "Normalisation classifications niv8, nettoyage libellés"],
        ["hospitals.py", "etablissements_hospitaliers.csv", "hospitals_paris_clean.csv", "Extraction coordonnées GPS, Paris uniquement"],
        ["green_spaces.py", "espaces_verts.geojson", "espaces_verts_paris.geojson", "Validation géométries, extraction surface réelle"],
        ["transport.py", "arrets.csv + velib.csv", "transport_arrets_paris.csv + velib_paris.csv", "Normalisation types, filtrage Paris"],
        ["thermal_comfort.py", "les-arbres.geojson + ilots-de-fraicheur.geojson", "thermal_comfort_base.geojson", "Agrégation arbres + îlots par zone IRIS"],
        ["sale_price.py", "PDF historique prix", "sale_price_m2.csv", "Extraction tables PDF via pdfplumber, parsing prix"],
        ["rent_price.py", "Base_OP_2024 + KML + XLSX", "rent_data_complet.csv", "Jointure zones KML + données loyers, conversion virgule→point"],
    ]
)

heading2("4.3 Couche Gold — scripts de calcul des indicateurs")
body(
    "La couche Gold est le cœur analytique du projet. Chaque script Gold lit les CSV Silver, "
    "effectue des calculs géospatiaux (buffer, jointure spatiale, distance haversine), "
    "normalise les résultats sur 0–10 et écrit les résultats en CSV + MySQL."
)

body("Technologie clé : GeoPandas avec projection Lambert-93 (EPSG:2154)")
body(
    "Pourquoi Lambert-93 ? Les calculs de distance (buffer en mètres) nécessitent un "
    "système de coordonnées en mètres. Le GRS80 (WGS84) donne des degrés, pas des mètres. "
    "Lambert-93 est la projection officielle de l'IGN pour la France métropolitaine."
)

code_block(
    "# Exemple de buffer spatial (utilisé pour tous les indicateurs)\n"
    "iris_lambert = iris_gdf.to_crs(epsg=2154)              # Projection métrique\n"
    "buffered = iris_lambert.copy()\n"
    "buffered['geometry'] = iris_lambert.buffer(500)         # 500 m autour de chaque zone\n"
    "join = gpd.sjoin(schools_lambert, buffered, predicate='within')  # Jointure spatiale\n"
    "counts = join.groupby('IRIS')['school'].count()         # Comptage"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. LES 4 INDICATEURS — ALGORITHMES DÉTAILLÉS
# ══════════════════════════════════════════════════════════════════════════════

heading1("5. Les 4 indicateurs personnalisés — algorithmes détaillés")

body(
    "Le projet produit 4 indicateurs composites originaux. Chacun fusionne plusieurs "
    "sources de données via des algorithmes géospatiaux sur mesure."
)

# ── 5.1 VIVABILITÉ FAMILIALE ───────────────────────────────────────────────────
heading2("5.1 Indice de Vivabilité Familiale")

body(
    "C'est l'indicateur phare du projet. Il combine 5 piliers en un score composite 0–10 "
    "représentant la qualité de vie globale pour les familles dans chaque zone IRIS."
)

heading3("Formule composite")
code_block(
    "vivabilite_score =  0,20 × école_score\n"
    "                  + 0,20 × santé_score\n"
    "                  + 0,20 × transport_score\n"
    "                  + 0,20 × services_quotidiens_score\n"
    "                  + 0,20 × espaces_verts_score"
)

body("Chaque pilier est normalisé indépendamment sur 0–10 avant d'être agrégé.")

heading3("Pilier 1 — Accessibilité aux écoles")
bullet("Données : 3 fichiers Excel (collèges, maternelles, élémentaires) → 1 377 établissements fusionnés")
bullet("Algorithme : buffer spatial de 500 m autour du centroïde de chaque zone IRIS")
bullet("Métrique : nombre d'écoles dans le buffer / population × 1 000 = écoles pour 1 000 habitants")
bullet("Normalisation : min-max sur l'ensemble des zones parisiennes → 0–10")

heading3("Pilier 2 — Accès aux soins de santé")
bullet("Données : hôpitaux (pondération × 3,0) + établissements médicaux BDCOM (pharmacies, médecins — pondération × 1,5)")
bullet("Algorithme : buffer 500 m, comptage pondéré")
bullet("weighted_healthcare_access = nb_hôpitaux × 3,0 + nb_services_médicaux × 1,5")
bullet("Normalisation min-max → 0–10")
bullet("Justification des poids : un hôpital vaut 2× une pharmacie en termes d'accès aux soins")

heading3("Pilier 3 — Transport (composante Vivabilité)")
bullet("Données : arrêts IDFM (métro, RER, tram, bus, câbleway) + stations Vélib'")
bullet("Poids par type : métro=1,0 · RER=1,2 · tram=0,7 · bus=0,4 · câbleway=0,5 · vélib=0,3")
bullet("Algorithme : buffer 500 m, somme des poids des arrêts dans le buffer")
bullet("weighted_stops = Σ(poids_type × nb_arrêts_type)")
bullet("Normalisation min-max → 0–10")
bullet("Justification : le RER a le plus grand poids car il couvre des distances bien plus importantes")

heading3("Pilier 4 — Services quotidiens")
bullet("Données : BDCOM 2023, filtré sur les catégories niv8 : alimentaire (2), service commercial (4)")
bullet("Inclut : épiceries, banques, bureaux de poste, bibliothèques, équipements sportifs, culture")
bullet("Exclut : santé (traité séparément dans le pilier 2)")
bullet("Algorithme : buffer 500 m, comptage des établissements (poids = 1,0 par établissement)")
bullet("Normalisation min-max → 0–10")

heading3("Pilier 5 — Espaces verts accessibles")
bullet("Données : espaces_verts_paris.geojson (parcs, jardins, squares avec surface réelle en m²)")
bullet("Algorithme en 2 passes :")
bullet("   Pass 1 — Espaces INTÉRIEURS : centroïde de l'espace vert dans la zone IRIS → score 100%", level=1)
bullet("   Pass 2 — Espaces ADJACENTS : centroïde dans un buffer de 300 m autour de la zone → score 50%", level=1)
bullet("total_vert_m2 = m2_intérieur + 0,5 × m2_adjacent")
bullet("vert_par_habitant = total_vert_m2 / population")
bullet("Normalisation min-max → 0–10")
bullet("Justification du centroïde : évite le double-comptage des grands parcs (Bois de Boulogne) qui chevauchent plusieurs zones")
bullet("Justification du bonus 50% : un habitant peut marcher jusqu'à un parc voisin — il mérite un crédit partiel")

heading3("Gestion des données manquantes")
body(
    "Si une zone IRIS manque un sous-score (jointure spatiale vide), la valeur est remplacée "
    "par la médiane de toutes les zones de Paris pour ce pilier. "
    "Cela évite de pénaliser les zones à faible densité tout en maintenant une cohérence statistique."
)

heading3("Résultat final")
add_table(
    ["Métrique", "Valeur"],
    [
        ["Score minimum", "~2,5/10"],
        ["Score maximum", "~8,5/10"],
        ["Score moyen", "~6,2/10"],
        ["Zone #1 (rang 1)", "Montparnasse 5 — Paris 14e — 6,2/10 (après repondération verts)"],
        ["Classement", "1 = meilleure zone, N = pire zone"],
    ]
)

page_break()

# ── 5.2 TRANSPORT ─────────────────────────────────────────────────────────────
heading2("5.2 Indicateur d'Accessibilité aux Transports")

body(
    "Cet indicateur est distinct du pilier transport de la vivabilité. Il utilise une "
    "méthodologie différente (distance haversine + décroissance) et une échelle 0–1."
)

heading3("Algorithme — Distance Haversine")
code_block(
    "Pour chaque zone IRIS :\n"
    "  1. Calculer le centroïde de la zone (lat, lng)\n"
    "  2. Pour chaque arrêt de transport dans un rayon de 800 m :\n"
    "       d = haversine(centroïde_zone, arrêt)      # en mètres\n"
    "       sum_poids += poids_type\n"
    "       sum_pondéré_distance += poids_type × d\n"
    "  3. density_raw = log(1 + sum_poids)             # logarithmique\n"
    "  4. distance_moy_pondérée = sum_pondéré / sum_poids\n"
    "\n"
    "  5. Score densité    = normalisation(density_raw)      [0–1]\n"
    "  6. Score proximité  = 1 - (distance_moy / 800)        [0–1, clampé]\n"
    "\n"
    "  7. transport_score = 0,60 × score_densité + 0,40 × score_proximité"
)

bullet("Rayon : 800 m (vs 500 m pour la vivabilité — représente ~10 min à pied)")
bullet("Pourquoi log ? La relation arrêts–confort est sous-linéaire : passer de 2 à 10 arrêts apporte moins de valeur que passer de 0 à 2")
bullet("Pourquoi 60/40 densité/proximité ? La densité (variété des lignes) prime sur la simple proximité")

page_break()

# ── 5.3 CONFORT THERMIQUE ─────────────────────────────────────────────────────
heading2("5.3 Indice de Confort Thermique Urbain")

body(
    "Cet indicateur mesure la capacité d'une zone IRIS à offrir un refuge face à la "
    "chaleur urbaine (îlots de chaleur). Il combine deux dimensions complémentaires."
)

heading3("Formule")
code_block(
    "densite_arbres  = nb_arbres / (surface_iris_m2 / 10 000)   [arbres/hectare]\n"
    "ratio_fraicheur = surface_ilots_fraicheur / surface_iris    [% de la zone]\n"
    "\n"
    "score_arbres    = normalisation_minmax(densite_arbres) × 100\n"
    "score_fraicheur = normalisation_minmax(ratio_fraicheur) × 100\n"
    "\n"
    "indice_confort_thermique = 0,40 × score_arbres + 0,60 × score_fraicheur"
)

bullet("Échelle : 0–100 (affiché 0–10 sur la carte en divisant par 10)")
bullet("Poids fraîcheur 60% : les îlots de fraîcheur sont plus efficaces thermiquement que la densité d'arbres seule")
bullet("Cas d'usage : identifier les quartiers à risque lors d'épisodes caniculaires")
bullet("Format de sortie : Parquet (préserve la géométrie pour Geopandas → export direct en GeoJSON)")

page_break()

# ── 5.4 LOGEMENT ──────────────────────────────────────────────────────────────
heading2("5.4 Indicateur Logement — Prix de vente et loyers")

body(
    "Contrairement aux autres indicateurs (échelle IRIS), les données de prix immobiliers "
    "ne sont disponibles qu'à l'échelle de l'arrondissement (20 zones). Deux sous-indicateurs :"
)

heading3("Prix de vente (DVF + PDF historique)")
bullet("Source 1 : DVF 2025 — transactions immobilières brutes (filtrage : Appart/Maison, 5–1000 m², 1k–50k €/m²)")
bullet("Source 2 : PDF historique de la Mairie de Paris (extraction pdfplumber)")
bullet("Agrégation : médiane des prix/m² par arrondissement")
bullet("Normalisation : higher_is_better=False → les arrondissements les moins chers obtiennent le score le plus élevé")
bullet("Format : Parquet + MySQL")

heading3("Loyers (Observatoire des Loyers Paris 2024)")
bullet("Source : Base_OP_2024_L7501.csv — loyers par zone élémentaire")
bullet("Agrégation par arrondissement : loyer_median_m2, loyer_q1_m2, loyer_q3_m2")
bullet("Les Q1 et Q3 permettent d'afficher la distribution (l'écart-type du marché)")
bullet("Normalisation : higher_is_better=False → loyer élevé = mauvais score")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. BASES DE DONNÉES
# ══════════════════════════════════════════════════════════════════════════════

heading1("6. Base de données — MySQL (Gold) et MongoDB (Analytics)")

heading2("6.1 MySQL — Stockage des indicateurs Gold")
body(
    "MySQL stocke les résultats calculés par le pipeline. L'API lit ces tables au démarrage "
    "et les charge en mémoire (DataStore). Les requêtes MySQL ne sont donc pas sur le chemin critique."
)

add_table(
    ["Table MySQL", "Colonnes principales", "Utilisation"],
    [
        ["school_density", "IRIS, school_count, schools_per_1000, school_score, population", "Indicateur écoles / API /indicators/schools"],
        ["transport_score", "IRIS, stop_count, weighted_stops, transport_score", "Composante transport pour vivabilité"],
        ["transport_score_iris", "CODE_IRIS, density_score, proximity_score, transport_score", "API /indicators/transport (échelle 0–1)"],
        ["transport_points", "id, name, type, lat, lng", "Table de référence — points de transport sur la carte"],
        ["services_score", "IRIS, hospital_count, service_count, weighted_services", "Composante services pour vivabilité"],
        ["green_spaces_score", "IRIS, interior_m2, adjacent_m2, total_green_m2, green_m2_per_resident", "Composante espaces verts"],
        ["healthcare_score", "IRIS, hospital_count, healthcare_service_count, weighted_healthcare_access", "Composante santé pour vivabilité"],
        ["daily_services_score", "IRIS, daily_service_count, weighted_daily_service_count", "Composante services quotidiens"],
        ["vivabilite_familiale", "IRIS, school_score, healthcare_score, transport_score, daily_services_score, green_spaces_score, vivabilite_score, vivabilite_rank", "Indicateur composite principal"],
        ["thermal_comfort", "code_iris, densite_arbres, ratio_fraicheur, indice_confort_thermique", "Indicateur confort thermique"],
        ["sale_price_median", "arrondissement, prix_m2, geometry", "Prix de vente par arrondissement"],
        ["rent_data", "c_ar, loyer_median_m2, loyer_q1_m2, loyer_q3_m2", "Loyers par arrondissement"],
    ]
)

heading2("6.2 MongoDB Atlas — Authentification et Analytics")
body(
    "MongoDB est utilisé pour deux cas d'usage différents, qui ne nécessitent pas de schéma rigide "
    "(NoSQL adapté) :"
)

heading3("Collection 'users' — Authentification")
code_block(
    "{\n"
    "  _id: ObjectId,\n"
    "  email: String (unique, indexé),\n"
    "  username: String (unique, indexé),\n"
    "  hashed_password: String (bcrypt + SHA-256),\n"
    "  is_active: Boolean,\n"
    "  created_at: DateTime,\n"
    "  updated_at: DateTime\n"
    "}"
)

heading3("Collection 'zone_user_interests' — Analytics")
code_block(
    "{\n"
    "  _id: ObjectId,\n"
    "  user_key: String (UUID anonyme),\n"
    "  zone_id: String (code IRIS ou 'arr-XX'),\n"
    "  zone_name: String,\n"
    "  geography: String ('iris' | 'arrondissement'),\n"
    "  clicks: Integer (incrémenté à chaque clic),\n"
    "  created_at: DateTime,\n"
    "  updated_at: DateTime\n"
    "}\n"
    "Indexes : (user_key, zone_id) unique · zone_id (pour les agrégats)"
)

body(
    "Chaque fois qu'un utilisateur clique sur une zone de la carte, une requête "
    "fire-and-forget est envoyée à /analytics/zone-clicks. Cela permet de voir "
    "quelles zones sont les plus consultées, sans impact sur les performances de la carte."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. API REST — FASTAPI
# ══════════════════════════════════════════════════════════════════════════════

heading1("7. API REST — FastAPI")

heading2("7.1 Démarrage et chargement des données")
body(
    "FastAPI utilise un gestionnaire de contexte de vie (lifespan) pour charger toutes les "
    "données au démarrage :"
)
code_block(
    "@asynccontextmanager\n"
    "async def lifespan(app: FastAPI):\n"
    "    app.state.data = DataStore.load()  # Chargement unique en mémoire\n"
    "    yield"
)

body(
    "Le DataStore charge en mémoire : toutes les CSVs Gold, les GeoJSON (IRIS + arrondissements), "
    "les scores vivabilité/transport/thermique/loyers. Chaque requête HTTP lit directement "
    "depuis la RAM — pas de requête SQL en temps réel."
)

heading2("7.2 Endpoints disponibles")
add_table(
    ["Préfixe", "Exemple de route", "Description"],
    [
        ["/auth", "POST /auth/register, /auth/login, GET /auth/me", "Inscription, connexion JWT, profil utilisateur"],
        ["/map", "GET /map/vivabilite-familiale", "GeoJSON IRIS avec scores vivabilité (fond de carte)"],
        ["/map", "GET /map/vivabilite-familiale/arrondissement", "GeoJSON arrondissement (zoom dézoomé)"],
        ["/map", "GET /map/thermal-comfort", "GeoJSON IRIS avec indice de confort thermique"],
        ["/map", "GET /map/housing/rent", "GeoJSON arrondissement avec loyers médians"],
        ["/map", "GET /map/housing/sale", "GeoJSON arrondissement avec prix de vente"],
        ["/indicators/vivabilite-familiale", "GET / · /{code_iris} · /arrondissements", "Scores vivabilité, filtrage, pagination"],
        ["/indicators/transport", "GET / · /{code_iris} · /points", "Scores transport + points d'arrêt"],
        ["/indicators/schools", "GET / · /{code_iris} · /arrondissements", "Densité scolaire par zone"],
        ["/iris", "GET / · /{code_iris}", "Métadonnées IRIS (nom, code, superficie)"],
        ["/analytics/zone-clicks", "POST (enregistrer) · GET /zones/top", "Suivi des clics sur les zones (MongoDB)"],
        ["/health", "GET /health", "Nombre de lignes chargées par dataset"],
        ["/stats", "GET /stats", "Statistiques agrégées globales"],
    ]
)

heading2("7.3 CORS et sécurité")
bullet("CORS configuré pour localhost:3000 (dev) et domaine de production")
bullet("Routes de carte publiques (pas d'authentification requise — données publiques)")
bullet("Routes analytics authentifiées en option mais non bloquantes")
bullet("Validation Pydantic sur toutes les entrées (email, mot de passe min 8 chars, etc.)")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. FRONTEND
# ══════════════════════════════════════════════════════════════════════════════

heading1("8. Frontend — Carte interactive (Next.js + Mapbox GL)")

heading2("8.1 Architecture des pages")
add_table(
    ["Route", "Page", "Accès", "Description"],
    [
        ["/", "page.tsx", "Public", "Landing page avec présentation des indicateurs et statistiques"],
        ["/login", "login/page.tsx", "Public", "Formulaire de connexion (email + mot de passe)"],
        ["/register", "register/page.tsx", "Public", "Formulaire d'inscription (username + email + mot de passe)"],
        ["/dashboard", "dashboard/page.tsx", "Protégé (AuthGuard)", "Carte interactive complète avec tous les indicateurs"],
    ]
)

heading2("8.2 Composant principal — EnhancedMapDashboard")
body(
    "Ce composant React (2 000+ lignes) est le cœur de l'expérience utilisateur. "
    "Il gère :"
)
bullet("4 indicateurs principaux (Vivabilité, Transport, Confort Thermique, Logement)")
bullet("Changement de source de données selon le niveau de zoom (zoom < 11 → arrondissements, zoom ≥ 11 → IRIS)")
bullet("Calcul en temps réel du score composite avec poids ajustables par l'utilisateur")
bullet("Classement dynamique des zones selon le score actif")
bullet("Sidebar avec recherche, filtres par arrondissement, score minimum")
bullet("Popup de détail au clic/survol sur chaque zone")
bullet("Enregistrement analytics des clics (fire-and-forget vers MongoDB)")

heading2("8.3 Logique de zoom adaptatif")
code_block(
    "const ZOOM_BREAK = 11;  // Seuil de bascule\n"
    "\n"
    "const sourceData = useMemo(() => {\n"
    "  if (mainIndicator === 'thermal') return thermalData;\n"
    "  if (mainIndicator === 'housing') return selectedMetric === 'sale_score' ? saleData : rentData;\n"
    "  // Vivabilité & Transport : zoom adaptatif\n"
    "  if (currentZoom < ZOOM_BREAK && arrData) return arrData;  // 20 arrondissements\n"
    "  return vivabiliteData;  // 800+ zones IRIS\n"
    "}, [mainIndicator, currentZoom, arrData, vivabiliteData, ...]);"
)
body(
    "Quand l'utilisateur dézoome sous le niveau 11, la carte bascule automatiquement "
    "sur les polygones d'arrondissement (données agrégées). Un badge dans la carte "
    "indique le niveau actif et invite à zoomer pour plus de détail."
)

heading2("8.4 Calcul interactif des poids")
code_block(
    "// L'utilisateur peut ajuster les poids des 5 piliers en temps réel\n"
    "function calculateWeightedScore(properties, weights) {\n"
    "  let weightedSum = 0, availableWeight = 0;\n"
    "  for (const pillar of pillarMeta) {\n"
    "    const value = properties[pillar.key];\n"
    "    const weight = weights[pillar.key];\n"
    "    if (typeof value === 'number' && weight > 0) {\n"
    "      weightedSum += value * weight;\n"
    "      availableWeight += weight;\n"
    "    }\n"
    "  }\n"
    "  return availableWeight > 0 ? weightedSum / availableWeight : null;\n"
    "}"
)
body(
    "Ce calcul est effectué entièrement côté client (en mémoire, dans le navigateur). "
    "Modifier les poids recalcule instantanément les scores, les couleurs de la carte "
    "et le classement — sans aucun appel API supplémentaire."
)

heading2("8.5 Coloration de la carte (Mapbox GL expressions)")
code_block(
    '// Expression Mapbox pour la couleur de remplissage\n'
    '"fill-color": [\n'
    '  "case",\n'
    '  ["!=", ["typeof", ["get", "active_score"]], "number"],\n'
    '  "#94a3b8",  // Gris pour les zones sans données\n'
    '  ["interpolate", ["linear"], ["get", "active_score"],\n'
    '    0, "#f43f5e",   // Rouge   (mauvais)\n'
    '    2, "#fb923c",   // Orange\n'
    '    4, "#fbbf24",   // Jaune\n'
    '    6, "#a3e635",   // Vert clair\n'
    '    8, "#22c55e",   // Vert\n'
    '   10, "#047857"    // Vert foncé (excellent)\n'
    '  ]\n'
    ']'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. AUTHENTIFICATION JWT
# ══════════════════════════════════════════════════════════════════════════════

heading1("9. Système d'authentification JWT")

heading2("9.1 Flux d'authentification complet")
code_block(
    "INSCRIPTION\n"
    "  1. Formulaire /register → POST /auth/register\n"
    "  2. Validation Pydantic (email valide, username 3-32 chars, mdp min 8 chars)\n"
    "  3. Hash du mot de passe : sha256(mdp) → bcrypt(sha256, salt)\n"
    "     [SHA-256 en pré-traitement pour dépasser la limite 72 octets de bcrypt]\n"
    "  4. Insertion MongoDB (index unique sur email + username)\n"
    "  5. Retour : UserOut (id, email, username, created_at)\n"
    "\n"
    "CONNEXION\n"
    "  1. Formulaire /login → POST /auth/login\n"
    "  2. Récupération utilisateur par email depuis MongoDB\n"
    "  3. Vérification : bcrypt.checkpw(sha256(mdp_fourni), hash_stocké)\n"
    "  4. Génération JWT HS256 : { sub: user_id, exp: now + 30 min }\n"
    "  5. Retour : { access_token, token_type: 'bearer' }\n"
    "  6. Stockage côté client : localStorage['ude_auth_token']\n"
    "\n"
    "ACCÈS PROTÉGÉ\n"
    "  1. AuthGuard (React) vérifie localStorage au montage\n"
    "  2. Si pas de token → redirect /login\n"
    "  3. Si token présent → GET /auth/me avec header Authorization: Bearer <token>\n"
    "  4. FastAPI décode le JWT, vérifie la signature, retourne UserOut"
)

heading2("9.2 Choix techniques de sécurité")
bullet("Pas de passlib : incompatible avec bcrypt 5.x (bug detect_wrap_bug). Implémentation directe avec la lib bcrypt.")
bullet("SHA-256 pré-hash : contourne la limite de 72 octets de bcrypt pour les mots de passe longs")
bullet("JWT HS256 avec SECRET_KEY de 32 octets (définie en variable d'environnement)")
bullet("Données sensibles dans .env (MONGO_URI, SECRET_KEY) — non committées")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════

heading1("10. Analytics utilisateurs — MongoDB")

body(
    "Chaque clic sur une zone de la carte est enregistré de manière anonyme. "
    "Un UUID est généré et stocké dans localStorage la première fois qu'un utilisateur "
    "visite la carte — pas de cookies, pas d'email requis."
)

heading2("Fonctionnement")
code_block(
    "// lib/api.ts — côté frontend\n"
    "function recordZoneClick(payload) {\n"
    "  fetch('/analytics/zone-clicks', {\n"
    "    method: 'POST',\n"
    "    body: JSON.stringify({ user_key, zone_id, zone_name, geography })\n"
    "  }).catch(() => {})  // fire-and-forget : les erreurs ne cassent pas la carte\n"
    "}"
)

heading2("Ce qu'on peut analyser")
bullet("Zones les plus consultées de tout Paris (classement global)")
bullet("Zones préférées d'un utilisateur spécifique (par son UUID)")
bullet("Comparaison entre les zones d'intérêt et les scores vivabilité")
bullet("Heatmap d'intérêt vs données objectives")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. STACK TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════

heading1("11. Stack technique complète")

add_table(
    ["Couche", "Technologie", "Version", "Rôle et justification"],
    [
        ["Pipeline données", "Python", "3.11+", "Standard data engineering, riche écosystème géospatial"],
        ["Pipeline données", "Pandas", "2.0+", "Manipulation de DataFrames tabulaires"],
        ["Pipeline données", "GeoPandas", "0.14+", "Opérations géospatiales (buffer, jointure, projection)"],
        ["Pipeline données", "Shapely", "2.0+", "Géométries vectorielles (polygones, points, buffers)"],
        ["Pipeline données", "pdfplumber", "0.11+", "Extraction de tables depuis PDF"],
        ["Pipeline données", "pyarrow", "14+", "Format Parquet (géométries + données numériques)"],
        ["API Backend", "FastAPI", "0.110+", "Framework async Python, OpenAPI auto-généré"],
        ["API Backend", "SQLAlchemy", "2.0+", "ORM pour MySQL"],
        ["API Backend", "PyMySQL", "1.1+", "Driver MySQL pur Python"],
        ["API Backend", "PyMongo", "4.6+", "Driver MongoDB (Atlas)"],
        ["API Backend", "python-jose", "3.3+", "JWT HS256"],
        ["API Backend", "bcrypt", "4.0+", "Hachage des mots de passe"],
        ["API Backend", "Pydantic", "2.0+", "Validation des modèles de données"],
        ["Base de données", "MySQL", "8.0", "Tables Gold (indicateurs calculés)"],
        ["Base de données", "MongoDB Atlas", "Cloud", "Utilisateurs + analytics clics"],
        ["Frontend", "Next.js", "16.2.4", "Framework React (SSR + App Router)"],
        ["Frontend", "React", "19.2.4", "Interface utilisateur réactive"],
        ["Frontend", "Mapbox GL JS", "3.x", "Rendu carte vectorielle WebGL haute performance"],
        ["Frontend", "react-map-gl", "8.x", "Wrapper React pour Mapbox"],
        ["Frontend", "Tailwind CSS", "v4", "CSS utilitaire — design system cohérent"],
        ["Frontend", "TypeScript", "5", "Typage statique du frontend"],
        ["Infrastructure", "Docker Compose", "-", "Stack locale (MySQL + API + Web)"],
    ]
)

heading2("Pourquoi ces choix ?")
body("Quelques justifications architecturales importantes :")
bullet("GeoPandas + Lambert-93 : la projection Lambert-93 est obligatoire pour faire des calculs de distance en mètres. Sans elle, les buffers seraient en degrés (incohérent).")
bullet("Données chargées en mémoire au démarrage : les GeoJSON IRIS font ~50 MB. Les charger une fois au démarrage plutôt qu'à chaque requête réduit la latence API de secondes à millisecondes.")
bullet("Format Parquet pour les géométries : les fichiers avec colonnes geometry (WKT/WKB) ne se sauvegardent pas proprement en CSV. Parquet préserve les types binaires.")
bullet("MongoDB pour l'authentification : schéma flexible (évolution sans migration) et hébergement Atlas gratuit avec 512 MB.")
bullet("MySQL pour les Gold : données tabulaires structurées avec jointures — SQL est ici le bon outil.")
bullet("Mapbox GL : rendu WebGL des tuiles vectorielles — capable d'afficher 800+ polygones colorés en temps réel à 60 FPS.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. CE QUI FONCTIONNE BIEN
# ══════════════════════════════════════════════════════════════════════════════

heading1("12. Ce qui fonctionne bien")

heading2("Pipeline de données")
bullet("Architecture Medallion propre et reproductible — tout développeur peut re-exécuter le pipeline depuis zéro")
bullet("Séparation stricte Bronze/Silver/Gold — aucun calcul métier dans les couches Silver")
bullet("Gestion gracieuse des fichiers manquants — l'API démarre même si un dataset est absent")
bullet("Poids VIVABILITE_WEIGHTS centralisés dans config.py — modification sans toucher au code métier")

heading2("API")
bullet("Réponses quasi-instantanées grâce au chargement en mémoire du DataStore au démarrage")
bullet("Endpoints GeoJSON prêts à consommer directement par Mapbox GL — zéro transformation côté frontend")
bullet("Documentation OpenAPI auto-générée accessible sur /docs")
bullet("CORS configuré proprement pour le développement local et la production")

heading2("Frontend — Carte interactive")
bullet("Zoom adaptatif : bascule automatique arrondissement ↔ IRIS sans rechargement des données")
bullet("Repondération temps réel : l'utilisateur ajuste les 5 piliers et voit les scores et le classement changer instantanément")
bullet("Zones sans données (parcs, zones industrielles) affichées en gris neutre — pas de trous blancs")
bullet("Page de démarrage publique avec landing page professionnelle")

heading2("Sécurité")
bullet("Mots de passe hachés avec bcrypt (salt aléatoire) + SHA-256 pré-hash")
bullet("JWT avec expiration (30 minutes) — pas de tokens permanents")
bullet("Variables sensibles dans .env — non commitées dans le dépôt")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. LIMITES ET AMÉLIORATIONS
# ══════════════════════════════════════════════════════════════════════════════

heading1("13. Limites actuelles et axes d'amélioration")

heading2("Limites techniques actuelles")
bullet("Données non déployées : le pipeline tourne en local, pas de CI/CD qui maintient les données à jour automatiquement")
bullet("Pas de données temporelles sur la carte : l'historique des prix existe en base mais pas de slider temporel dans l'interface")
bullet("Loyers et prix uniquement à l'échelle arrondissement : pas de granularité IRIS pour les données financières")
bullet("Pas de comparaison côte à côte : impossible de comparer deux arrondissements dans l'interface actuelle")
bullet("Pipeline Gold lent : les jointures spatiales GeoPandas sur 800+ zones × 5 000+ points peuvent prendre 10–20 minutes en local")
bullet("Chargement en mémoire à l'API : si de nombreux datasets sont ajoutés, la RAM de la machine hébergeant l'API deviendra un goulot d'étranglement")

heading2("Améliorations à court terme (1–2 mois)")
bullet("Déploiement sur serveur cloud (Railway, Render, ou AWS EC2) avec GitHub Actions pour recalcul périodique")
bullet("Slider temporel pour l'évolution des prix immobiliers (données DVF 2014–2025 disponibles)")
bullet("Mode comparaison : sélectionner 2 zones et voir un tableau comparatif côte à côte")
bullet("Export PDF du profil d'une zone (score détaillé + carte + graphiques)")

heading2("Améliorations à moyen terme (3–6 mois)")
bullet("Granularité prix à l'échelle IRIS : croiser DVF avec les codes IRIS via géocodage inverse")
bullet("Pipeline en streaming (Apache Kafka ou AWS Kinesis) pour les données de prix en temps réel")
bullet("Cache Redis pour les GeoJSON — évite de recharger les ~50 MB à chaque redémarrage API")
bullet("Tests d'intégration complets — la couverture actuelle est partielle")
bullet("Score d'accessibilité : intégrer les données de revenu médian pour calculer un taux d'effort (loyer/revenu)")

heading2("Améliorations à long terme (6+ mois)")
bullet("Intégration de nouvelles sources : qualité de l'air (Airparif), criminalité (stats.interieur.gouv.fr), crèches, déchetteries")
bullet("Moteur de recommandation personnalisé : 'Trouvez votre quartier idéal' selon profil famille")
bullet("Version mobile avec React Native ou PWA")
bullet("Dashboard B2B : API payante pour agences immobilières")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

heading1("14. Conclusion")

body(
    "Urban Data Explorer démontre qu'il est possible de construire une plateforme "
    "d'analyse urbaine complète — de la donnée brute jusqu'à la carte interactive — "
    "uniquement à partir de données publiques et de technologies open-source."
)

heading2("Ce que nous avons construit")
bullet("Un pipeline de données reproductible sur 12 sources hétérogènes (CSV, Excel, GeoJSON, PDF, KML, TXT)")
bullet("4 indicateurs composites originaux avec des algorithmes géospatiaux justifiés")
bullet("Une API REST FastAPI avec 15+ endpoints documentés et des réponses en millisecondes")
bullet("Une carte interactive professionnelle avec zoom adaptatif, repondération temps réel et analytics")
bullet("Un système d'authentification sécurisé (JWT + bcrypt + MongoDB)")

heading2("Ce que cela démontre techniquement")
bullet("Maîtrise du patron Medallion (Bronze/Silver/Gold) en data engineering")
bullet("Compétences géospatiales : projections Lambert-93, buffers, jointures spatiales, distance haversine")
bullet("Intégration full-stack : Python (pipeline) → FastAPI (API) → React/Mapbox (frontend)")
bullet("Architecture orientée performance : données en mémoire, calculs côté client, fire-and-forget analytics")

heading2("Message final")
p = doc.add_paragraph()
run = p.add_run(
    "Le territoire parisien n'est pas homogène. Les données le prouvent chaque jour. "
    "Urban Data Explorer rend ces inégalités visibles, mesurables et compréhensibles — "
    "pour que les choix de vie puissent être guidés par les données, pas par les préjugés."
)
run.italic = True
run.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

credits = doc.add_paragraph()
credits.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_c = credits.add_run("Projet EFREI 2024–2025  •  Urban Data Explorer  •  Data Engineering & Cartographie")
run_c.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# ── Sauvegarde ─────────────────────────────────────────────────────────────────
output_path = os.path.join(
    r"c:\Users\Adrien Duval\Documents\Efrei\Project-urban-data-explorer\Project-urban-data-explorer",
    "Urban_Data_Explorer_Presentation_EFREI.docx"
)
doc.save(output_path)
print(f"Document sauvegardé : {output_path}")
